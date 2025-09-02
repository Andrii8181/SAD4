from docx import Document
from datetime import datetime
import os
import platform
import subprocess

def export_results_to_word(indicator, units, data, result):
    filename = "Результати_аналізу.docx"

    doc = Document()
    doc.add_heading("SAD - Статистичний аналіз даних", level=0)

    doc.add_paragraph(f"Назва показника: {indicator}")
    doc.add_paragraph(f"Одиниці виміру: {units}")

    # Таблиця з початковими даними
    doc.add_heading("Початкові дані", level=1)
    table = doc.add_table(rows=1, cols=1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Значення"
    for value in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(value)

    # Результати тесту
    doc.add_heading("Перевірка нормальності (Шапіро-Вілк)", level=1)
    doc.add_paragraph(f"W = {result['W']:.4f}, p = {result['p']:.4f}")
    if result["normal"]:
        doc.add_paragraph("✅ Дані відповідають нормальному розподілу (p > 0.05).")
    else:
        doc.add_paragraph("⚠️ Дані НЕ відповідають нормальному розподілу (p ≤ 0.05).")

    # Дата аналізу
    doc.add_paragraph(f"\nДата виконання: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    doc.save(filename)

    # Автоматично відкриваємо файл
    open_file(filename)


def open_file(filepath):
    """Відкриває файл після створення у Word / LibreOffice"""
    if platform.system() == "Windows":
        os.startfile(filepath)
    elif platform.system() == "Darwin":  # macOS
        subprocess.call(["open", filepath])
    else:  # Linux
        subprocess.call(["xdg-open", filepath])
